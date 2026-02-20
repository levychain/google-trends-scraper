#!/usr/bin/env python3
"""Generate Word document: Why WW3/War Risk Is Our Next Vertical"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
import datetime

doc = Document()

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

# -- Title --
title = doc.add_heading('Why WW3/War Risk Is Our Next Vertical', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = subtitle.add_run(f'Internal Strategy Brief — {datetime.date.today().strftime("%B %d, %Y")}')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.italic = True

doc.add_paragraph()

# -- Executive Summary --
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'The same playbook that worked for Epstein Files — high-intent keywords, unmatched cadence, '
    'connected narrative, source-backed depth — is sitting wide open in the war/geopolitics space. '
    'No dominant WW3 podcast exists. The demand is confirmed by hard money. The window is open now.'
)

# -- The Opportunity --
doc.add_heading('The Opportunity: Hard Data', level=1)

doc.add_heading('Market Signals (as of February 2026)', level=2)

signals = [
    ('$402M+', 'traded on Polymarket across 62 active Iran conflict markets'),
    ('83%', 'probability of "US strikes Iran by Dec 31, 2026" on Polymarket (January peak)'),
    ('4%', 'single-day oil price jump after VP Vance said strikes are "on the table" (Feb 18)'),
    ('$91/bbl', 'BloombergNEF oil projection by Q4 2026 if Iran exports disrupted'),
    ('89 sec', 'Doomsday Clock reading — closest to midnight in history'),
    ('2,000%', 'spike in "World War III" searches in one week (June 2025, Yahoo News)'),
]

for metric, description in signals:
    p = doc.add_paragraph()
    run_bold = p.add_run(f'{metric}  ')
    run_bold.bold = True
    run_bold.font.size = Pt(12)
    p.add_run(f'— {description}')

doc.add_heading('Key Events This Week', level=2)

events = [
    'Iran partially closed the Strait of Hormuz (Feb 17, 2026)',
    'Oil surged on Vance "strikes on the table" statement (Feb 18, 2026)',
    'Geneva nuclear talks between U.S. and Iran ongoing',
    'Baba Vanga "WW3 in 2026" prediction viral on TikTok/X globally',
    'Selective Service auto-registration changed Dec 2025 — 18-26 year old males now auto-enrolled',
]

for event in events:
    doc.add_paragraph(event, style='List Bullet')

# -- Google Trends Comparison --
doc.add_heading('Google Trends: War Keywords vs. Epstein', level=1)

doc.add_paragraph(
    'We ran head-to-head Google Trends comparisons using our scraper tool, normalizing war keywords '
    'directly against "Jeffrey Epstein" over the past 12 months. Here is what the data shows.'
)

doc.add_heading('"Iran" vs "Jeffrey Epstein" — 12 Months, Same Normalization', level=2)

doc.add_paragraph(
    'When "iran" spikes, it is so massive that Epstein rounds to zero in the same normalization context. '
    'Iran\'s June 2025 peak was 12x bigger than Epstein\'s biggest week in this pairing.'
)

# Table 1: Iran vs Epstein
table1 = doc.add_table(rows=5, cols=3)
table1.style = 'Light Grid Accent 1'
table1.alignment = WD_TABLE_ALIGNMENT.LEFT

headers = ['Period', '"iran"', '"Jeffrey Epstein"']
for i, header in enumerate(headers):
    cell = table1.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

data1 = [
    ('Baseline (Mar–May 2025)', '1–2', '0'),
    ('Iran strike spike (Jun 2025)', '58–100', '0'),
    ('Post-spike crash (Jul–Oct 2025)', '1–5', '0–3'),
    ('Current elevated (Jan–Feb 2026)', '4–19', '1–8'),
]

for row_idx, (period, iran_val, epstein_val) in enumerate(data1, 1):
    table1.rows[row_idx].cells[0].text = period
    table1.rows[row_idx].cells[1].text = iran_val
    table1.rows[row_idx].cells[2].text = epstein_val

doc.add_paragraph()

doc.add_heading('"World War 3" vs "Jeffrey Epstein" — 12 Months, Same Normalization', level=2)

doc.add_paragraph(
    'Epstein is at its all-time 12-month high right now (100 in week of Feb 1, 2026), which inflates '
    'the comparison. Despite this, WW3\'s June 2025 spike (17–18) shows the keyword activates '
    'strongly during geopolitical events.'
)

# Table 2: WW3 vs Epstein
table2 = doc.add_table(rows=6, cols=3)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.LEFT

headers2 = ['Period', '"world war 3"', '"Jeffrey Epstein"']
for i, header in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

data2 = [
    ('Baseline (Mar–May 2025)', '1–2', '1–2'),
    ('WW3 spike (Jun 2025)', '17–18', '2'),
    ('Epstein spike (Jul 2025)', '1', '34'),
    ('Epstein spike (Nov 2025)', '1', '33'),
    ('Epstein spike (Feb 2026 — now)', '1', '100'),
]

for row_idx, (period, ww3_val, epstein_val) in enumerate(data2, 1):
    table2.rows[row_idx].cells[0].text = period
    table2.rows[row_idx].cells[1].text = ww3_val
    table2.rows[row_idx].cells[2].text = epstein_val

doc.add_paragraph()

# -- Cluster Analysis --
doc.add_heading('The Cluster Advantage', level=2)

doc.add_paragraph(
    'The comparison above is individual keywords. Our strategy doesn\'t target one keyword — '
    'it targets the entire cluster simultaneously. The combined search surface includes:'
)

cluster_keywords = [
    ('"iran"', 'Peaks at 100. Currently 4–19 (elevated).'),
    ('"world war 3" / "ww3"', 'Peaks at 18. Currently 1–2 (baseline).'),
    ('"iran war"', 'Currently 4 (elevated).'),
    ('"gas prices"', 'Always high volume. 55.3 similarity score to Epstein.'),
    ('"oil prices"', 'Always high volume. Spiking this week.'),
    ('"military draft"', 'Spikes with events. Legislation change driving searches.'),
    ('"strait of hormuz"', 'Spiking this week (Feb 17 partial closure).'),
    ('"baba vanga 2026"', 'Viral globally on TikTok/X right now.'),
    ('50+ long-tail variations', 'Safe countries, nuclear fallout, NATO Article 5, etc.'),
]

for keyword, note in cluster_keywords:
    p = doc.add_paragraph()
    run_bold = p.add_run(f'{keyword}  ')
    run_bold.bold = True
    p.add_run(f'— {note}')

doc.add_paragraph(
    'The combined search surface of the war keyword cluster exceeds Epstein\'s single-topic volume '
    'during geopolitical escalation periods — which we are currently in.'
)

# -- Pattern Comparison --
doc.add_heading('Pattern Comparison: Epstein vs. War Topic', level=1)

# Table 3: Pattern comparison
table3 = doc.add_table(rows=7, cols=3)
table3.style = 'Light Grid Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.LEFT

headers3 = ['Metric', 'Epstein Files', 'War/WW3 Cluster']
for i, header in enumerate(headers3):
    cell = table3.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

data3 = [
    ('Search pattern', 'Steady baseline + event spikes', 'Spike-driven + elevated baseline during crises'),
    ('Peak volume', 'High (100 in current window)', 'Higher ceiling ("iran" dwarfs Epstein at peak)'),
    ('Baseline volume', 'Moderate (sustained by content)', 'Lower individually, larger as cluster'),
    ('Keyword cluster size', '~20 core terms', '60+ core terms across 6 arcs'),
    ('Competitive landscape', 'Low competition (owned it)', 'No dominant WW3 podcast. Name unclaimed.'),
    ('Demand driver', 'Document releases, hearings, podcast', 'Geopolitical events, policy, viral predictions'),
]

for row_idx, (metric, epstein, war) in enumerate(data3, 1):
    table3.rows[row_idx].cells[0].text = metric
    table3.rows[row_idx].cells[1].text = epstein
    table3.rows[row_idx].cells[2].text = war

doc.add_paragraph()

# -- Competitive Gap --
doc.add_heading('Competitive Gap: Why Nobody Owns This', level=1)

doc.add_paragraph(
    'There is no dominant WW3-specific podcast. The name is literally unclaimed — "WW3 Podcast" '
    'on Spotify is about sports and college life.'
)

doc.add_heading('Institutional players (high credibility, low cadence)', level=2)
institutional = [
    'War on the Rocks — 313 episodes over 12 years. Think-tank audience. Dry.',
    'GZERO World (Ian Bremmer) — Weekly geopolitics. Mainstream but low-cadence.',
    'Pod Save the World — Former Obama NSC staff. Political framing.',
    'CSIS / Foreign Affairs podcasts — Policy-nerd audience. Weekly.',
]
for item in institutional:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('The gap we exploit', level=2)
doc.add_paragraph(
    'These institutional podcasts are thorough but dry. They don\'t speak to the person Googling '
    '"am I going to get drafted" or "what happens to gas prices if Iran closes Hormuz." '
    'That person — scared, confused, seeking clarity — is our audience. '
    'Nobody is serving them with volume, accessibility, and connected narrative.'
)

# -- The Plan --
doc.add_heading('The Plan', level=1)

plan_items = [
    ('60 time-capsule episodes', 'dropped at once as a connected baseline narrative. Same model we proved with 64 Epstein episodes.'),
    ('2 episodes/day ongoing', 'live feed tracking escalation in real time. AM update + PM analysis.'),
    ('6 narrative arcs', 'Foundations → Military-Intel → Flashpoints → U.S. Decision Room → Futures → Global Picture.'),
    ('Source-backed', 'Tiered sourcing model with confidence scoring. Same evidence standard as Epstein Files.'),
    ('SEO-first titles', 'Every title optimized for click-through with specific claims, numbers, and personal-stakes hooks.'),
    ('Web property', 'Topic Explorer-style site with transcripts, enrichment data, and connected episode graph for long-tail SEO.'),
]

for title_text, description in plan_items:
    p = doc.add_paragraph()
    run_bold = p.add_run(f'{title_text}  ')
    run_bold.bold = True
    p.add_run(f'— {description}')

# -- Sample Enhanced Titles --
doc.add_heading('Sample Episode Titles (Enhanced)', level=2)

sample_titles = [
    'Ep 1: How Close Are We to World War 3 Right Now',
    'Ep 5: How a Middle East War Becomes World War 3 in 72 Hours',
    'Ep 12: Can Iron Dome Stop 1,000 Missiles at Once? The Math Says No',
    'Ep 21: What $9 Gas Looks Like If Iran Closes the Strait of Hormuz',
    'Ep 22: You Were Auto-Registered for the Draft in December — Here\'s What It Means',
    'Ep 35: The U.S. Is Running Out of Missiles and Can\'t Make Them Fast Enough',
    'Ep 38: The Fake News That Could Start World War 3',
    'Ep 51: The 10 Safest Countries If World War 3 Breaks Out',
    'Ep 52: Baba Vanga Predicted World War 3 in 2026 — Here\'s What the Evidence Says',
    'Ep 60: $402 Million in Bets Say War Is Coming — What Prediction Markets Know',
]

for t in sample_titles:
    doc.add_paragraph(t, style='List Bullet')

# -- Bottom Line --
doc.add_heading('Bottom Line', level=1)

p = doc.add_paragraph()
run = p.add_run(
    'The war/WW3 keyword cluster has a higher ceiling than Epstein (iran at peak dwarfs Epstein to zero) '
    'but a lower floor (spike-driven, crashes between events). The bet is that the current geopolitical '
    'environment — Iran nuclear standoff, Hormuz closure, $402M in prediction market bets, viral WW3 '
    'predictions — keeps the floor elevated long enough to build the audience. '
    'Nobody owns this space. We have the playbook. The window is open.'
)

doc.add_paragraph()
p_final = doc.add_paragraph()
run_final = p_final.add_run('The window is open. It won\'t stay open forever.')
run_final.bold = True
run_final.font.size = Pt(13)

# -- Save --
output_path = Path('/Users/adamlevy/Google Trends Scraper/output/reports/War_Vertical_Strategy_Brief.docx')
doc.save(str(output_path))
print(f"Document saved to: {output_path}")
